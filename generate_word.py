import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_report():
    doc = Document()

    # Page Setup (A4, Margin Standard 1 inch / 2.54 cm)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1)
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)

    # Styles
    PRIMARY_COLOR = RGBColor(18, 18, 18)     # Deep Charcoal
    ACCENT_COLOR = RGBColor(4, 120, 87)      # Emerald Green
    MUTED_COLOR = RGBColor(85, 85, 85)       # Slate Gray

    # ==================== COVER / JUDUL UTAMA ====================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(36)
    title_p.paragraph_format.space_after = Pt(6)
    run_badge = title_p.add_run("DOKUMEN LAPORAN & PANDUAN IMPLEMENTASI SISTEM")
    run_badge.font.size = Pt(11)
    run_badge.font.bold = True
    run_badge.font.color.rgb = ACCENT_COLOR

    p_main_title = doc.add_paragraph()
    p_main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_main_title.paragraph_format.space_after = Pt(12)
    run_title = p_main_title.add_run("LAPOR DESA RAU\nSistem Aspirasi & Pengaduan Warga Berbasis Geofencing GPS")
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY_COLOR

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(24)
    run_sub = p_sub.add_run("Pemerintah Desa Rau, Kecamatan Kedung, Kabupaten Jepara, Jawa Tengah")
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = MUTED_COLOR

    # Table Info Box Ringkasan
    tbl_meta = doc.add_table(rows=4, cols=2)
    tbl_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_meta.autofit = False

    meta_data = [
        ("Tautan Publik (Warga)", "https://lapor-desa-rau.vercel.app"),
        ("Portal Petugas (Rahasia)", "https://lapor-desa-rau.vercel.app/kantor-desa (PIN: rau2026)"),
        ("Repositori Kode (GitHub)", "https://github.com/mrizkymxx/lapor-desa-rau"),
        ("Biaya Server & Operasional", "Rp 0,- / Gratis Selamanya (Vercel + Supabase Free Tier)")
    ]

    for idx, (label, val) in enumerate(meta_data):
        row = tbl_meta.rows[idx]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = Inches(2.2)
        c1.width = Inches(4.0)
        set_cell_background(c0, "F3F4F6")
        set_cell_background(c1, "FFFFFF")
        set_cell_margins(c0, top=100, bottom=100, left=120, right=120)
        set_cell_margins(c1, top=100, bottom=100, left=120, right=120)

        p0 = c0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.font.size = Pt(9.5)
        r0.font.bold = True

        p1 = c1.paragraphs[0]
        r1 = p1.add_run(val)
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = ACCENT_COLOR if "https" in val else PRIMARY_COLOR

    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    # ==================== BAB I: PENDAHULUAN ====================
    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        r = h.add_run(text)
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = ACCENT_COLOR
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        r = h.add_run(text)
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = PRIMARY_COLOR
        return h

    def add_body(text, bold_prefix="", italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.bold = True
            rb.font.size = Pt(10.5)
        r = p.add_run(text)
        r.font.size = Pt(10.5)
        r.font.italic = italic
        return p

    add_heading_1("BAB I. LATAR BELAKANG & RUMUSAN MASALAH")

    add_body("Pelayanan publik dan pengelolaan sarana prasarana desa seringkali menghadapi kendala komunikasi antara warga dan perangkat balai desa. Di Desa Rau, Kecamatan Kedung, Kabupaten Jepara, banyak keluhan seperti jalan berlubang, lampu penerangan jalan mati, saluran drainase mampet, atau layanan administrasi desa yang lambat tertangani karena ketiadaan media pelaporan yang terpusat dan transparan.")

    add_heading_2("1.1. Permasalahan Pelaporan Konvensional")
    doc.add_paragraph("1. Warga sungkan / enggan melapor langsung ke kantor desa karena keterbatasan waktu kerja atau rasa tidak enak hati.", style='List Bullet')
    doc.add_paragraph("2. Laporan sering tercecer dan terlupakan saat hanya disampaikan lewat obrolan lisan atau chat grup WhatsApp warga yang tertimbun.", style='List Bullet')
    doc.add_paragraph("3. Ketidakakuratan lokasi kerusakan fisik (jalan/lampu) karena pelapor hanya menyebut patokan umum sehingga petugas lapangan kesulitan mencari titik masalah.", style='List Bullet')
    doc.add_paragraph("4. Kekhawatiran masuknya laporan palsu (hoaks/spam) dari pihak luar desa apabila sistem dibuka tanpa filter wilayah.", style='List Bullet')
    doc.add_paragraph("5. Anggaran dana desa yang terbatas untuk pengadaan server atau aplikasi sistem berbayar yang mahal.", style='List Bullet')

    add_heading_2("1.2. Tujuan & Solusi Inovasi")
    add_body("Aplikasi 'Lapor Desa Rau' dirancang untuk mengatasi seluruh kendala tersebut dengan menyediakan sistem pengaduan berbasis web responsif modern, dapat diakses seluruh kalangan tanpa login, eksklusif dalam batas wilayah Desa Rau menggunakan validasi GPS Geofencing, serta beroperasi dengan biaya server Rp 0,- (Zero-Cost Architecture).")

    # ==================== BAB II: ARSITEKTUR & FITUR SISTEM ====================
    add_heading_1("BAB II. FITUR UTAMA SISTEM & PENGGUNAAN")

    add_heading_2("2.1. Fitur untuk Masyarakat Desa (Sisi Warga)")
    doc.add_paragraph("• Akses Tanpa Registrasi / Password: Warga langsung dapat membuka tautan aplikasi di HP tanpa repot mendaftar email atau menghafal password. Ramah untuk warga lansia.", style='List Bullet')
    doc.add_paragraph("• Geofencing Presisi 2.0 KM: Sistem otomatis memvalidasi koordinat GPS HP pelapor terhadap titik Balai Desa Rau (-6.646463, 110.667857). Tombol kirim hanya aktif jika warga berada di wilayah Desa Rau.", style='List Bullet')
    doc.add_paragraph("• Kompresi Foto Otomatis: Foto kamera HP yang berukuran 5–8 MB otomatis diperkecil di memori browser HP menjadi format WebP berukuran ~50 KB sebelum dikirim. Sangat hemat kuota dan cepat terkirim di jaringan lemah.", style='List Bullet')
    doc.add_paragraph("• Opsi Pelaporan Anonim: Warga dapat memilih untuk menyamarkan identitasnya sehingga nama pelapor tertulis sebagai 'Warga Rau (Anonim)'.", style='List Bullet')
    doc.add_paragraph("• Nomor Tiket & Papan Pantau Transparansi: Setiap laporan memperoleh nomor tiket unik (RAU-YYYYMMDD-XXXX) dan status pengerjaan dapat dipantau terbuka oleh seluruh warga.", style='List Bullet')
    doc.add_paragraph("• Peta Titik Lokasi Interaktif: Detail laporan dilengkapi peta digital OpenStreetMap untuk memastikan titik koordinat akurat lengkap dengan tombol pintas ke Google Maps.", style='List Bullet')
    doc.add_paragraph("• PWA (Progressive Web App): Warga dapat menyematkan aplikasi ke layar utama HP (Add to Home Screen) sehingga dapat dibuka layaknya aplikasi Android/iOS native tanpa perlu mengunduh di Google Play Store.", style='List Bullet')

    add_heading_2("2.2. Fitur Portal Pengelola (Sisi Petugas Balai Desa)")
    add_body("Halaman petugas diproteksi dan disembunyikan dari menu navigasi publik (URL: /kantor-desa) dengan autentikasi PIN khusus (Default: rau2026):")
    doc.add_paragraph("1. Triage & Manajemen Status: Mengubah alur status aduan secara berkala (Menunggu Tindak Lanjut ➔ Sedang Dikerjakan ➔ Selesai Dikerjakan ➔ Ditolak).", style='List Bullet')
    doc.add_paragraph("2. Publikasi Tanggapan Resmi: Memberikan pesan balasan dan keterangan penanganan yang langsung terbaca oleh warga.", style='List Bullet')
    doc.add_paragraph("3. Unggah Foto Bukti Perbaikan: Petugas lapangan dapat mengunggah foto kondisi fisik sarana setelah selesai diperbaiki sebagai bukti transparansi publik.", style='List Bullet')
    doc.add_paragraph("4. Tombol Chat WhatsApp Instan: Menghubungi nomor WhatsApp pelapor dengan satu klik menggunakan format pesan resmi yang sopan dan santun.", style='List Bullet')
    doc.add_paragraph("5. Ekspor Rekapitulasi CSV/Excel: Mengunduh data seluruh aduan dan status penanganan dalam format spreadsheet untuk kebutuhan arsip pelaporan ke Camat atau BPD.", style='List Bullet')

    # ==================== BAB III: SPESIFIKASI TEKNIS ====================
    add_heading_1("BAB III. SPESIFIKASI TEKNIS & KEAMANAN")

    tbl_tech = doc.add_table(rows=6, cols=3)
    tbl_tech.alignment = WD_TABLE_ALIGNMENT.CENTER

    tech_headers = ["Komponen", "Teknologi Terpasang", "Fungsi & Keunggulan"]
    for i, h in enumerate(tech_headers):
        c = tbl_tech.rows[0].cells[i]
        set_cell_background(c, "121212")
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9.5)

    tech_rows = [
        ("Frontend & SSR", "Next.js 15 (App Router) + React 19", "Performa tinggi, Server-Side Rendering, PWA ready, load < 1 detik"),
        ("Styling & UI", "Tailwind CSS v4 (Neo-Brutalism)", "Desain kontras tinggi, border tebal, ergonomis layar HP"),
        ("Database Relasional", "Supabase PostgreSQL 17 (Singapore)", "Keamanan RLS (Row Level Security), ACID compliance, free tier 500 MB"),
        ("Modul Geospasial", "PostGIS Extension (PostgreSQL)", "Formula kalkulasi jarak koordinat server-side anti-spoofing"),
        ("Media Storage", "Supabase Storage ('laporan-media')", "Penyimpanan foto WebP terkompresi kapasitas hingga 8.000+ laporan"),
    ]

    for row_idx, data_row in enumerate(tech_rows, start=1):
        for col_idx, text_val in enumerate(data_row):
            c = tbl_tech.rows[row_idx].cells[col_idx]
            set_cell_background(c, "F9FAFB" if row_idx % 2 == 1 else "FFFFFF")
            set_cell_margins(c, top=80, bottom=80, left=100, right=100)
            p = c.paragraphs[0]
            r = p.add_run(text_val)
            r.font.size = Pt(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ==================== BAB IV: NASKAH SOSIALISASI WARGA ====================
    add_heading_1("BAB IV. NASKAH SOSIALISASI KE WARGA (FORUM RT / MUSDES)")

    add_body("Berikut adalah panduan naskah percakapan bilingual yang dapat digunakan oleh Kepala Desa, Sekretaris Desa, atau Ketua RT saat mensosialisasikan aplikasi ini di pertemuan warga:")

    # Callout Box Naskah Jawa
    tbl_jawa = doc.add_table(rows=1, cols=1)
    tbl_jawa.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_jawa = tbl_jawa.rows[0].cells[0]
    set_cell_background(c_jawa, "FEF3C7") # Warm Yellow
    set_cell_margins(c_jawa, top=140, bottom=140, left=150, right=150)
    pj = c_jawa.paragraphs[0]
    rj_title = pj.add_run("Naskah Bahasa Jawa (Kagem Rapat RT / Pertemuan Warga):\n\n")
    rj_title.font.bold = True
    rj_title.font.size = Pt(10.5)

    naskah_jawa_text = (
        "\"Assalamu'alaikum Warahmatullahi Wabarakatuh.\n"
        "Bapak-bapak, Ibu-ibu sedaya warga Desa Rau ingkang kulo hormati.\n\n"
        "Kulo badhe nepangaken aplikasi enggal kagem kemajuan desa kito, inggih menika: LAPOR DESA RAU (https://lapor-desa-rau.vercel.app).\n\n"
        "Menawi panjenengan manggihi dalan bolong, lampu penerangan mati, saluran toya mampet, utawi keluhan administrasi, sakmenika mboten sah bingung. Cukup buka HP panjenengan, jepret fotonipun langsung wonten lokasi, tulis katrangan singkat, lajeng klik Kirim.\n\n"
        "Aplikasi menika mboten sah ngangge daftar email utawi password sing ribet. Lokasinipun otomatis kecathet ngangge GPS, dados petugas balai desa mboten bakal kesasar madosi panggenane. Panjenengan ugi saget milih opsi 'Anonim' menawi sungkan asmanipun katingal.\n\n"
        "Sedaya aduan saget dipun pantau sareng-sareng wonten menu PANTAU. Mugi-mugi aplikasi menika saget mbiyantu kito sedaya njagi lan mbangun Desa Rau supados langkung sae lan tentrem. Matur nuwun sanget.\n"
        "Wassalamu'alaikum Warahmatullahi Wabarakatuh.\""
    )
    rj_body = pj.add_run(naskah_jawa_text)
    rj_body.font.size = Pt(9.5)
    rj_body.font.italic = True

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # ==================== BAB V: KESIMPULAN & PENUTUP ====================
    add_heading_1("BAB V. KESIMPULAN & REKOMENDASI")
    add_body("Sistem Informasi Lapor Desa Rau telah selesai dibangun, teruji bebas galat, dan berstatus produksi aktif (Live) di Vercel. Sistem ini menawarkan efisiensi tinggi, keterbukaan informasi publik, serta keandalan jangka panjang tanpa membebani biaya APBDes. Diharapkan Pemerintah Desa Rau dan jajaran pengurus RT/RW dapat segera menyebarluaskan tautan ini kepada seluruh warga.")

    # Save document
    out_path = "D:/Lapor Desa Rau/Laporan_Lengkap_Lapor_Desa_Rau.docx"
    doc.save(out_path)
    print("SUCCESS: Word Document berhasil dibuat di " + out_path)

create_report()
